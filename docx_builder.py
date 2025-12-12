"""
DOCX Builder Module

Based on json results from OCR engine, build a DOCX file.
json results provide:
+ blocks with each block has document entities label like 'paragraph', 'heading', 'list', 'table', 'image', etc.
+ content (text or table data) for each block
+ bbox + style info for each block
=> Build a DOCX file with proper structure and formatting. (formatting doesn't need to be perfect, can be whatever)

"""

from docx import Document
from docx.shared import Pt # pt: points for font size
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # for paragraph alignment
import json
import os

from config import Config

# block_label appearance so far:
# 'text', 'doc_title', 'paragraph_title', 'table', 'image', 'number' (likely page number)

# Style rule for each block_label:
# 'doc_title': Title, centered, larger font size
# 'paragraph_title': Heading 1, centered, slightly larger font size
# 'text': Normal paragraph style


class DOCXBuilder:
    def __init__(self):
        """
        Initialize the DOCX Builder.
        """
        self.document = Document()
        self._setup_default_styles()

    def _setup_default_styles(self):
        """
        Setup default styles for the document.
        """
        styles = self.document.styles

        # doc_title
        title_style = styles['Title']  # 1: Paragraph style
        title_font = title_style.font
        title_font.size = Pt(24)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # paragraph_title
        heading1_style = styles['Heading 1']  # 1: Paragraph style
        heading1_font = heading1_style.font
        heading1_font.size = Pt(18)
        heading1_font.bold = True
        heading1_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # text
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.size = Pt(12)
        normal_font.bold = False


    def add_title(self, text):
        """
        Add a title to the document.

        :param text: Text content of the title.
        """
        title_paragraph = self.document.add_paragraph(text, style='Title')

    def add_heading(self, text, level=1):
        """
        Add a heading to the document.

        :param text: Text content of the heading.
        :param level: Heading level (1-9).
        """
        heading_paragraph = self.document.add_heading(text, level=level)

    def add_paragraph(self, text, style=None, alignment=None):
        """
        Add a paragraph to the document.

        :param text: Text content of the paragraph.
        :param style: Optional style for the paragraph.
        :param alignment: Optional alignment for the paragraph.
        """
        paragraph = self.document.add_paragraph(text)
       
    def add_image(self, image_path, width=None, height=None):
        """
        Add an image to the document.

        :param image_path: Path to the image file.
        :param width: Optional width for the image.
        :param height: Optional height for the image.
        """

        # image that is ocr'ed get saved as separate files, so we can just add them here

        if width and height:
            self.document.add_picture(image_path, width=width, height=height)
        elif width:
            self.document.add_picture(image_path, width=width)
        elif height:
            self.document.add_picture(image_path, height=height)
        else:
            self.document.add_picture(image_path)


    # add_table
    def add_table(self, table_html):
        """
        Add a table to the document from HTML format.
        
        :param table_html: HTML string containing table data
        """
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(table_html, 'html.parser')

        print("Parsing table HTML to add to DOCX...")
        table_tag = soup.find('table')
        if not table_tag:
            return  # No table found
        
        rows = table_tag.find_all('tr')
        if not rows:
            return
        
        # Calculate max columns (considering colspan)
        num_cols = 0
        for row in rows:
            cols_in_row = 0
            for cell in row.find_all(['td', 'th']):
                colspan = int(cell.get('colspan', 1))
                cols_in_row += colspan
            num_cols = max(num_cols, cols_in_row)
        
        num_rows = len(rows)
        
        # Create table
        table = self.document.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light List Accent 1'
        
        # Populate table
        for r, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            col_index = 0
            
            for cell in cells:
                colspan = int(cell.get('colspan', 1))
                text = cell.get_text(strip=True)
                
                # Handle merged cells
                if colspan > 1:
                    # Merge cells horizontally
                    merged_cell = table.cell(r, col_index).merge(table.cell(r, col_index + colspan - 1))
                    merged_cell.text = text
                else:
                    table.cell(r, col_index).text = text
                
                # Apply bold formatting for header cells (th tags or second row typically)
                is_header = cell.name == 'th' or (r == 1 and len(cells) > 1)
                if is_header:
                    for paragraph in table.cell(r, col_index).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                col_index += colspan

    def add_page_number(self, page_number):
        """
        Add a page number to the document.

        :param page_number: Page number to add.
        """
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(str(page_number))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    
    def add_page_break(self):
        """
        Add a page break to the document.
        """
        self.document.add_page_break()

    def save(self, file_path):
        """
        Save the document to the specified file path.

        :param file_path: Path to save the DOCX file.
        """
        self.document.save(file_path)

def build_docx_from_ocr_json(res_path, save_path):
    """
    Build a DOCX file from OCR JSON results.

    :param ocr_json: OCR results in JSON format.
    :param save_path: Path to save the generated DOCX file.
    """
    docx_builder = DOCXBuilder()

    # res_path has:
    # [page_number]_res.json
    # img folder with images
    # image name rule: img_in_image_box_block["block_bbox"][0]_box["block_bbox"][1]_box["block_bbox"][2]_box["block_bbox"][3].jpg
    # for example: block_bbox = [34, 45, 200, 300] -> image name: img_in_image_box_34_45_200_300.jpg
    # There is also an image for table img_in_table_box_... -> ignore for now

    # example of ocr json: .\\output\620\620_res.json
    # res_path = .\\output\620\
    
    page_number = os.path.basename(res_path.rstrip(os.sep))
    ocr_json = os.path.join(res_path, f"{page_number}_res.json")

    ocr_improved_json = ocr_json.replace("_res.json", "_improved.json")

    if os.path.exists(ocr_improved_json):
        ocr_json = ocr_improved_json

    with open(ocr_json, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    for block in data.get('parsing_res_list', []):
        label = block.get('block_label', 'text')
        content = block.get('block_content', '')

        if label == 'doc_title':
            docx_builder.add_title(content)
        elif label == 'paragraph_title':
            docx_builder.add_heading(content, level=1)
        elif label == 'text':
            docx_builder.add_paragraph(content)
        elif label == 'image':

            # images are in imgs/
            image_path = os.path.join(res_path, "imgs")
            # find image path based on block bbox
            block_bbox = block.get('block_bbox', [])
            if len(block_bbox) == 4:
                image_path = os.path.join(image_path, "img_in_image_box_" + \
                    f"{block_bbox[0]}_{block_bbox[1]}_{block_bbox[2]}_{block_bbox[3]}.jpg")

            if image_path:
                docx_builder.add_image(image_path)
                
        elif label == 'table':
            docx_builder.add_table(content)
        elif label == 'number':
            # 'number' is only a page number if:
            # + its bbox is the final in json
            # + its block_content is only a number from x to xxx
            if content.strip().isdigit() and block == data.get('parsing_res_list', [])[-1]:
                docx_builder.add_page_number(content.strip())
        

    docx_builder.save(save_path)